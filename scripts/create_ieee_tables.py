import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)

def set_two_columns(section):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '708') # 0.5 inch spacing

set_two_columns(doc.sections[0])

def add_table(doc, num, title, headers, rows):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(2)
    run1 = p1.add_run(num)
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(10)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(title)
    run2.font.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(10)
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            run = row_cells[i].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() 

add_table(doc, "TABLE I", "CLASSICAL MODEL COMPARISON", 
          ["Model", "Accuracy", "Macro F1"],
          [["Logistic Regression", "0.897", "0.893"],
           ["Random Forest", "0.945", "0.944"],
           ["XGBoost (16 feature)", "0.969", "0.968"],
           ["XGBoost (10 feature)", "0.969", "0.968"]])

add_table(doc, "TABLE II", "ABLATION STUDY RESULTS",
          ["Model Configuration", "Circuit Type", "Output Head", "Accuracy", "Macro F1"],
          [["Single Output (Baseline)", "6-qubit Basic", "expval(Z₀)", "0.3286", "0.2044"],
           ["Multi Output Basic", "6-qubit Basic", "6 expvals", "0.4019", "0.3393"],
           ["Multi Output Strongly", "6-qubit Strongly", "6 expvals", "0.4965", "0.4285"]])

add_table(doc, "TABLE III", "VQC TRAINING PROGRESS",
          ["Epoch", "Train Loss", "Test Accuracy", "Macro F1"],
          [["10", "1.4025", "0.4303", "0.3649"],
           ["20", "1.3141", "0.4563", "0.3998"],
           ["40", "1.2087", "0.4799", "0.4247"],
           ["60", "1.1471", "0.5272", "0.4852"],
           ["80", "1.1118", "0.5390", "0.5003"],
           ["100 (Best)", "1.0896", "0.5674", "0.5304"]])

add_table(doc, "TABLE IV", "CLASS-WISE METRICS",
          ["Class", "Precision", "Recall", "F1", "Support"],
          [["Insufficient_Weight", "0.46", "0.70", "0.55", "54"],
           ["Normal_Weight", "0.36", "0.21", "0.26", "58"],
           ["Overweight_Level_I", "0.55", "0.21", "0.30", "58"],
           ["Overweight_Level_II", "0.46", "0.43", "0.45", "58"],
           ["Obesity_Type_I", "0.46", "0.47", "0.46", "70"],
           ["Obesity_Type_II", "0.67", "0.93", "0.78", "60"],
           ["Obesity_Type_III", "0.84", "0.98", "0.91", "65"],
           ["Macro Avg", "0.54", "0.56", "0.53", "423"],
           ["Weighted Avg", "0.55", "0.57", "0.54", "423"]])

# Ensure results dir exists
os.makedirs("results", exist_ok=True)
doc.save("results/IEEE_Tables_100_Epochs.docx")
